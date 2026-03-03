#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
章节分析器 - 识别论文章节、提取内容、分析类型
"""

import re
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from enum import Enum

try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx")
    exit(1)


class ChapterType(Enum):
    """章节类型"""
    ABSTRACT = "abstract"
    INTRODUCTION = "introduction"
    RELATED_WORK = "related_work"
    METHOD = "method"
    EXPERIMENT = "experiment"
    RESULT = "result"
    DISCUSSION = "discussion"
    CONCLUSION = "conclusion"
    OTHER = "other"


@dataclass
class Chapter:
    """章节数据结构"""
    name: str
    type: ChapterType
    start_index: int
    end_index: int
    content: str
    page_count: int = 0

    def get_word_count(self) -> int:
        """获取字数"""
        return len(self.content.split())

    def get_estimated_citations(self) -> int:
        """估计引用数量"""
        # 根据章节类型和内容长度估计引用数
        type_multipliers = {
            ChapterType.ABSTRACT: 0.5,
            ChapterType.INTRODUCTION: 2.0,
            ChapterType.RELATED_WORK: 1.5,
            ChapterType.METHOD: 0.3,
            ChapterType.EXPERIMENT: 0.2,
            ChapterType.RESULT: 0.2,
            ChapterType.DISCUSSION: 0.5,
            ChapterType.CONCLUSION: 0.1,
            ChapterType.OTHER: 0.2,
        }
        multiplier = type_multipliers.get(self.type, 0.2)
        word_count = self.get_word_count()
        return max(1, int(word_count / 100 * multiplier))


class ChapterAnalyzer:
    """章节分析器"""

    # 章节名称匹配模式
    CHAPTER_PATTERNS = {
        ChapterType.ABSTRACT: [r"摘要|abstract", r"^摘\s*要$"],
        ChapterType.INTRODUCTION: [r"引言|introduction|前言", r"^引\s*言$"],
        ChapterType.RELATED_WORK: [r"相关工作|related work|文献综述|literature review", r"^相关工作$"],
        ChapterType.METHOD: [r"方法|method|methodology|研究方法", r"^方法$"],
        ChapterType.EXPERIMENT: [r"实验|experiment|evaluation", r"^实验$"],
        ChapterType.RESULT: [r"结果|result|results", r"^结果$"],
        ChapterType.DISCUSSION: [r"讨论|discussion", r"^讨论$"],
        ChapterType.CONCLUSION: [r"结论|conclusion|总结", r"^结论$"],
    }

    def __init__(self):
        self.chapters: List[Chapter] = []

    def analyze_document(self, doc_path: str) -> List[Chapter]:
        """
        分析Word文档，识别章节

        Args:
            doc_path: Word文档路径

        Returns:
            章节列表
        """
        try:
            doc = Document(doc_path)
            self.chapters = self._extract_chapters(doc)
            return self.chapters
        except Exception as e:
            print(f"Error analyzing document: {e}")
            return []

    def _extract_chapters(self, doc: Document) -> List[Chapter]:
        """从文档中提取章节"""
        chapters = []
        current_chapter = None
        current_content = []

        for para_index, para in enumerate(doc.paragraphs):
            # 检测标题
            heading_level = self._detect_heading_level(para)

            if heading_level == 1:  # 一级标题 = 章节
                # 保存前一个章节
                if current_chapter:
                    current_chapter.content = "\n".join(current_content)
                    current_chapter.end_index = para_index - 1
                    chapters.append(current_chapter)

                # 创建新章节
                chapter_name = para.text.strip()
                chapter_type = self._classify_chapter_type(chapter_name)
                current_chapter = Chapter(
                    name=chapter_name,
                    type=chapter_type,
                    start_index=para_index,
                    end_index=para_index,
                    content=""
                )
                current_content = []
            else:
                # 添加到当前章节内容
                if current_chapter and para.text.strip():
                    current_content.append(para.text.strip())

        # 保存最后一个章节
        if current_chapter:
            current_chapter.content = "\n".join(current_content)
            current_chapter.end_index = len(doc.paragraphs) - 1
            chapters.append(current_chapter)

        return chapters

    def _detect_heading_level(self, para) -> Optional[int]:
        """检测段落的标题级别"""
        # 检查style
        if para.style.name.startswith('Heading'):
            try:
                return int(para.style.name.split()[-1])
            except:
                pass

        # 检查格式特征
        if para.runs:
            first_run = para.runs[0]
            if first_run.font.bold and first_run.font.size:
                size = first_run.font.size.pt
                if size >= 16:
                    return 1
                elif size >= 14:
                    return 2
                elif size >= 12:
                    return 3

        return None

    def _classify_chapter_type(self, chapter_name: str) -> ChapterType:
        """分类章节类型"""
        chapter_name_lower = chapter_name.lower()

        for chapter_type, patterns in self.CHAPTER_PATTERNS.items():
            for pattern in patterns:
                if re.search(pattern, chapter_name_lower, re.IGNORECASE):
                    return chapter_type

        return ChapterType.OTHER

    def get_chapter_summary(self) -> Dict:
        """获取章节摘要"""
        summary = {
            "total_chapters": len(self.chapters),
            "chapters": []
        }

        for chapter in self.chapters:
            summary["chapters"].append({
                "name": chapter.name,
                "type": chapter.type.value,
                "word_count": chapter.get_word_count(),
                "estimated_citations": chapter.get_estimated_citations()
            })

        return summary

    def get_chapters_by_type(self, chapter_type: ChapterType) -> List[Chapter]:
        """按类型获取章节"""
        return [ch for ch in self.chapters if ch.type == chapter_type]

    def get_search_priority(self) -> List[Tuple[Chapter, int]]:
        """
        获取搜索优先级

        Returns:
            [(chapter, priority), ...] 优先级越低越先处理
        """
        priority_map = {
            ChapterType.INTRODUCTION: 1,
            ChapterType.RELATED_WORK: 2,
            ChapterType.DISCUSSION: 3,
            ChapterType.METHOD: 4,
            ChapterType.EXPERIMENT: 5,
            ChapterType.RESULT: 6,
            ChapterType.CONCLUSION: 7,
            ChapterType.ABSTRACT: 8,
            ChapterType.OTHER: 9,
        }

        chapters_with_priority = [
            (ch, priority_map.get(ch.type, 10))
            for ch in self.chapters
        ]

        return sorted(chapters_with_priority, key=lambda x: x[1])


def main():
    """测试"""
    analyzer = ChapterAnalyzer()
    # 这里可以添加测试代码
    pass


if __name__ == "__main__":
    main()
