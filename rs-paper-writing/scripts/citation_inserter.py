#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
引用插入器 - 识别引用位置、插入引用、生成BibTeX、更新参考文献
"""

import re
from typing import List, Dict, Optional, Tuple
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx")
    exit(1)


class CitationInserter:
    """引用插入器"""

    def __init__(self):
        self.citations = {}  # {key: bibtex_entry}

    def identify_cite_positions(self, text: str) -> List[Tuple[int, int, str]]:
        """
        识别引用位置

        Args:
            text: 文本内容

        Returns:
            [(start, end, context), ...] 引用位置列表
        """
        positions = []

        # 查找[CITE]占位符
        for match in re.finditer(r'\[CITE\]', text):
            start = max(0, match.start() - 100)
            end = min(len(text), match.end() + 100)
            context = text[start:end]
            positions.append((match.start(), match.end(), context))

        return positions

    def insert_citation(self, doc: Document, position: int, citation_key: str) -> bool:
        """
        在文档中插入引用

        Args:
            doc: Word文档对象
            position: 插入位置（段落索引）
            citation_key: 引用键

        Returns:
            是否成功
        """
        try:
            if position < len(doc.paragraphs):
                para = doc.paragraphs[position]
                # 替换[CITE]为\cite{key}
                for run in para.runs:
                    if "[CITE]" in run.text:
                        run.text = run.text.replace("[CITE]", f"~\\cite{{{citation_key}}}")
                return True
            return False
        except Exception as e:
            print(f"Error inserting citation: {e}")
            return False

    def generate_bibtex(self, reference: Dict) -> str:
        """
        生成BibTeX条目

        Args:
            reference: 文献信息

        Returns:
            BibTeX字符串
        """
        # 生成引用键
        authors = reference.get("authors", [])
        first_author = authors[0].split()[-1] if authors else "Unknown"
        year = reference.get("year", "0000")
        key = f"{first_author}{year}"

        # 生成BibTeX
        title = reference.get("title", "")
        venue = reference.get("venue", "")
        year = reference.get("year", "")
        doi = reference.get("doi", "")

        bibtex = f"""@article{{{key},
  title={{{title}}},
  author={{{', '.join(authors[:3])}}},
  journal={{{venue}}},
  year={{{year}}}"""

        if doi:
            bibtex += f",\n  doi={{{doi}}}"

        bibtex += "\n}"

        return bibtex

    def update_bibliography(self, doc: Document, citations: Dict[str, str]) -> bool:
        """
        更新参考文献列表

        Args:
            doc: Word文档对象
            citations: {key: bibtex_entry}

        Returns:
            是否成功
        """
        try:
            # 查找或创建参考文献部分
            bib_para = None
            for para in doc.paragraphs:
                if "参考文献" in para.text or "References" in para.text:
                    bib_para = para
                    break

            if not bib_para:
                # 在文档末尾添加参考文献部分
                bib_para = doc.add_paragraph()
                bib_para.text = "参考文献"

            # 添加BibTeX条目
            for key, bibtex in citations.items():
                bib_para = doc.add_paragraph()
                bib_para.text = bibtex

            return True
        except Exception as e:
            print(f"Error updating bibliography: {e}")
            return False

    def export_bibtex(self, citations: Dict[str, str], output_file: str) -> bool:
        """
        导出BibTeX文件

        Args:
            citations: {key: bibtex_entry}
            output_file: 输出文件路径

        Returns:
            是否成功
        """
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                for key, bibtex in citations.items():
                    f.write(bibtex)
                    f.write("\n\n")
            return True
        except Exception as e:
            print(f"Error exporting bibtex: {e}")
            return False


def main():
    """测试"""
    inserter = CitationInserter()
    # 这里可以添加测试代码
    pass


if __name__ == "__main__":
    main()
