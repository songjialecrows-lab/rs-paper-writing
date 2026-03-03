#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word格式自动调整工具 - RS-Paper-Writing Skill
支持多个学校模板、自定义格式规则、批量处理、格式预览
"""

import json
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict
from enum import Enum

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx")
    exit(1)


class HeadingLevel(Enum):
    """标题级别"""
    H1 = 1
    H2 = 2
    H3 = 3


@dataclass
class FontStyle:
    """字体样式"""
    name: str = "宋体"
    size: int = 12
    bold: bool = False
    italic: bool = False
    color: Optional[str] = None

    def apply_to_run(self, run):
        """应用字体样式到run对象"""
        run.font.name = self.name
        run.font.size = Pt(self.size)
        run.font.bold = self.bold
        run.font.italic = self.italic
        if self.color:
            run.font.color.rgb = RGBColor(*self._parse_color(self.color))

    @staticmethod
    def _parse_color(color_str: str) -> Tuple[int, int, int]:
        """解析颜色字符串 (hex或rgb)"""
        if color_str.startswith('#'):
            hex_color = color_str.lstrip('#')
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        return (0, 0, 0)


@dataclass
class ParagraphStyle:
    """段落样式"""
    font: FontStyle = None
    alignment: str = "left"  # left, center, right, justify
    line_spacing: float = 1.5
    space_before: int = 0
    space_after: int = 0
    indent_first: float = 0.5  # 首行缩进（英寸）

    def __post_init__(self):
        if self.font is None:
            self.font = FontStyle()


@dataclass
class HeadingStyle:
    """标题样式"""
    level: int
    font: FontStyle = None
    paragraph: ParagraphStyle = None

    def __post_init__(self):
        if self.font is None:
            self.font = FontStyle()
        if self.paragraph is None:
            self.paragraph = ParagraphStyle()


@dataclass
class TemplateConfig:
    """模板配置"""
    name: str
    school: str
    description: str

    # 页面设置
    page_width: float = 8.5  # 英寸
    page_height: float = 11.0
    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.25
    margin_right: float = 1.25

    # 正文样式
    body_font: FontStyle = None
    body_paragraph: ParagraphStyle = None

    # 标题样式
    heading_styles: Dict[int, HeadingStyle] = None

    # 特殊格式
    figure_caption_font: FontStyle = None
    table_caption_font: FontStyle = None
    reference_font: FontStyle = None

    # 页眉页脚
    header_text: Optional[str] = None
    footer_text: Optional[str] = None

    def __post_init__(self):
        if self.body_font is None:
            self.body_font = FontStyle(name="宋体", size=12)
        if self.body_paragraph is None:
            self.body_paragraph = ParagraphStyle()
        if self.heading_styles is None:
            self.heading_styles = self._default_heading_styles()
        if self.figure_caption_font is None:
            self.figure_caption_font = FontStyle(name="宋体", size=10)
        if self.table_caption_font is None:
            self.table_caption_font = FontStyle(name="宋体", size=10)
        if self.reference_font is None:
            self.reference_font = FontStyle(name="宋体", size=10)

    @staticmethod
    def _default_heading_styles() -> Dict[int, HeadingStyle]:
        """默认标题样式"""
        return {
            1: HeadingStyle(
                level=1,
                font=FontStyle(name="黑体", size=16, bold=True),
                paragraph=ParagraphStyle(alignment="center", space_before=12, space_after=12)
            ),
            2: HeadingStyle(
                level=2,
                font=FontStyle(name="黑体", size=14, bold=True),
                paragraph=ParagraphStyle(space_before=6, space_after=6)
            ),
            3: HeadingStyle(
                level=3,
                font=FontStyle(name="宋体", size=12, bold=True),
                paragraph=ParagraphStyle(space_before=0, space_after=0)
            ),
        }


class WordFormatter:
    """Word文档格式转换器"""

    def __init__(self, template_config: TemplateConfig):
        self.config = template_config

    def format_document(self, input_path: str, output_path: str) -> bool:
        """
        格式化Word文档

        Args:
            input_path: 输入文档路径
            output_path: 输出文档路径

        Returns:
            是否成功
        """
        try:
            doc = Document(input_path)

            # 应用页面设置
            self._apply_page_settings(doc)

            # 应用段落格式
            self._apply_paragraph_formats(doc)

            # 应用页眉页脚
            self._apply_headers_footers(doc)

            # 保存文档
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error formatting document: {e}")
            return False

    def _apply_page_settings(self, doc: Document):
        """应用页面设置"""
        sections = doc.sections
        for section in sections:
            section.page_width = Inches(self.config.page_width)
            section.page_height = Inches(self.config.page_height)
            section.top_margin = Inches(self.config.margin_top)
            section.bottom_margin = Inches(self.config.margin_bottom)
            section.left_margin = Inches(self.config.margin_left)
            section.right_margin = Inches(self.config.margin_right)

    def _apply_paragraph_formats(self, doc: Document):
        """应用段落格式"""
        for para in doc.paragraphs:
            # 检测标题级别
            heading_level = self._detect_heading_level(para)

            if heading_level:
                self._apply_heading_format(para, heading_level)
            else:
                self._apply_body_format(para)

    def _detect_heading_level(self, para) -> Optional[int]:
        """检测段落是否为标题及其级别"""
        # 检查style
        if para.style.name.startswith('Heading'):
            try:
                return int(para.style.name.split()[-1])
            except:
                pass

        # 检查格式特征（加粗、较大字号等）
        if para.runs:
            first_run = para.runs[0]
            if first_run.font.bold and first_run.font.size and first_run.font.size >= Pt(14):
                return 1
            elif first_run.font.bold and first_run.font.size and first_run.font.size >= Pt(12):
                return 2

        return None

    def _apply_heading_format(self, para, level: int):
        """应用标题格式"""
        if level not in self.config.heading_styles:
            return

        heading_style = self.config.heading_styles[level]

        # 应用字体
        for run in para.runs:
            heading_style.font.apply_to_run(run)

        # 应用段落格式
        self._apply_paragraph_style(para, heading_style.paragraph)

    def _apply_body_format(self, para):
        """应用正文格式"""
        for run in para.runs:
            self.config.body_font.apply_to_run(run)

        self._apply_paragraph_style(para, self.config.body_paragraph)

    def _apply_paragraph_style(self, para, style: ParagraphStyle):
        """应用段落样式"""
        # 对齐方式
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = alignment_map.get(style.alignment, WD_ALIGN_PARAGRAPH.LEFT)

        # 行距
        para.paragraph_format.line_spacing = style.line_spacing

        # 段前段后间距
        para.paragraph_format.space_before = Pt(style.space_before)
        para.paragraph_format.space_after = Pt(style.space_after)

        # 首行缩进
        para.paragraph_format.first_line_indent = Inches(style.indent_first)

    def _apply_headers_footers(self, doc: Document):
        """应用页眉页脚"""
        for section in doc.sections:
            if self.config.header_text:
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.text = self.config.header_text

            if self.config.footer_text:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.text = self.config.footer_text


class TemplateManager:
    """模板管理器"""

    def __init__(self, template_dir: str = "data/templates"):
        self.template_dir = Path(template_dir)
        self.template_dir.mkdir(parents=True, exist_ok=True)

    def save_template(self, config: TemplateConfig) -> bool:
        """保存模板配置"""
        try:
            template_file = self.template_dir / f"{config.name}.json"

            # 转换为可序列化的格式
            config_dict = self._config_to_dict(config)

            with open(template_file, 'w', encoding='utf-8') as f:
                json.dump(config_dict, f, ensure_ascii=False, indent=2)

            return True
        except Exception as e:
            print(f"Error saving template: {e}")
            return False

    def load_template(self, template_name: str) -> Optional[TemplateConfig]:
        """加载模板配置"""
        try:
            template_file = self.template_dir / f"{template_name}.json"

            if not template_file.exists():
                print(f"Template not found: {template_name}")
                return None

            with open(template_file, 'r', encoding='utf-8') as f:
                config_dict = json.load(f)

            return self._dict_to_config(config_dict)
        except Exception as e:
            print(f"Error loading template: {e}")
            return None

    def list_templates(self) -> List[str]:
        """列出所有可用模板"""
        templates = []
        for file in self.template_dir.glob("*.json"):
            templates.append(file.stem)
        return sorted(templates)

    @staticmethod
    def _config_to_dict(config: TemplateConfig) -> dict:
        """将TemplateConfig转换为字典"""
        return {
            "name": config.name,
            "school": config.school,
            "description": config.description,
            "page_width": config.page_width,
            "page_height": config.page_height,
            "margin_top": config.margin_top,
            "margin_bottom": config.margin_bottom,
            "margin_left": config.margin_left,
            "margin_right": config.margin_right,
            "body_font": asdict(config.body_font),
            "body_paragraph": asdict(config.body_paragraph),
            "figure_caption_font": asdict(config.figure_caption_font),
            "table_caption_font": asdict(config.table_caption_font),
            "reference_font": asdict(config.reference_font),
            "header_text": config.header_text,
            "footer_text": config.footer_text,
        }

    @staticmethod
    def _dict_to_config(config_dict: dict) -> TemplateConfig:
        """将字典转换为TemplateConfig"""
        return TemplateConfig(
            name=config_dict.get("name", ""),
            school=config_dict.get("school", ""),
            description=config_dict.get("description", ""),
            page_width=config_dict.get("page_width", 8.5),
            page_height=config_dict.get("page_height", 11.0),
            margin_top=config_dict.get("margin_top", 1.0),
            margin_bottom=config_dict.get("margin_bottom", 1.0),
            margin_left=config_dict.get("margin_left", 1.25),
            margin_right=config_dict.get("margin_right", 1.25),
            body_font=FontStyle(**config_dict.get("body_font", {})),
            body_paragraph=ParagraphStyle(**config_dict.get("body_paragraph", )),
            figure_caption_font=FontStyle(**config_dict.get("figure_caption_font", {})),
            table_caption_font=FontStyle(**config_dict.get("table_caption_font", {})),
            reference_font=FontStyle(**config_dict.get("reference_font", {})),
            header_text=config_dict.get("header_text"),
            footer_text=config_dict.get("footer_text"),
        )


def main():
    """主函数 - 使用示例"""
    # 创建模板管理器
    manager = TemplateManager()

    # 创建默认模板
    default_config = TemplateConfig(
        name="default",
        school="Default University",
        description="Default thesis template"
    )

    manager.save_template(default_config)

    # 列出所有模板
    print("Available templates:", manager.list_templates())


if __name__ == "__main__":
    main()
